import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_tos_docx():
    doc = docx.Document()

    # Impostazione margini pagina
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Titolo principale
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("CONTRATTO DI LICENZA D'USO SAAS,\nTERMINI E CONDIZIONI DI SERVIZIO E MANLEVA LEGALE AI")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(27, 54, 93) # Navy Blue

    # Sottotitolo Società
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("M PROJECT S.R.L. — PIATTAFORMA CONDOSMART")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(59, 130, 246) # Blue Accent

    # Data aggiornamento
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run("Ultimo aggiornamento: 23 Luglio 2026")
    run_date.font.name = 'Arial'
    run_date.font.size = Pt(10)
    run_date.font.italic = True
    run_date.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph() # Spaziatore

    # Avviso iniziale
    p_intro = doc.add_paragraph()
    run_intro = p_intro.add_run("Il presente documento costituisce un contratto giuridicamente vincolante tra M PROJECT S.R.L. ed il Cliente/Utente. Si prega di leggerlo attentamente prima di completare la registrazione o di utilizzare la piattaforma CondoFAST.")
    run_intro.font.name = 'Arial'
    run_intro.font.size = Pt(10.5)
    run_intro.font.italic = True

    sections_content = [
        ("1. DEFINIZIONI ED IDENTIFICAZIONE DEI CONTRAENTI", [
            ("Fornitore del Servizio: ", "M PROJECT S.R.L. (di seguito indicata come \"M PROJECT SRL\", \"CondoFAST\" o \"Fornitore\"), con sede legale in [INSERIRE SEDE LEGALE], P.IVA / C.F. [INSERIRE P.IVA / CF], PEC: [INSERIRE EMAIL PEC], Email Assistenza: supporto@condosmart.it."),
            ("Cliente / Utente: ", "L’Amministratore di Condominio (persona fisica o società di amministrazione condominiale) o lo Studio Professionale registrato sulla Piattaforma che utilizza il Servizio nell'esercizio della propria attività professionale o imprenditoriale (Trattamento B2B)."),
            ("Condominio: ", "L'ente di gestione o fabbricato amministrato dal Cliente ed inserito sulla Piattaforma."),
            ("Piattaforma / Servizio: ", "Il software gestionale in modalità Software-as-a-Service (SaaS) accessibile all'indirizzo web di CondoFAST, inclusivo delle funzionalità di gestione contabile, ripartizione millesimale, lettura documentale tramite Intelligenza Artificiale (OCR), estrazione dati, riconciliazione bancaria, generazione documenti PDF ed email marketing."),
            ("Funzionalità AI: ", "I moduli informatici basati su algoritmi di machine learning e modelli linguistici di grandi dimensioni (LLM) forniti da terze parti (es. Google Gemini, Anthropic Claude) integrati nel Servizio per l'estrazione automatica di dati da fatture, estratti conto o verbali.")
        ]),
        ("2. OGGETTO DEL CONTRATTO E AMBITO B2B", [
            ("", "1. M PROJECT SRL concede al Cliente una licenza d'uso temporanea, non esclusiva, non trasferibile e revocabile per l'accesso e l'utilizzo della Piattaforma CondoFAST via web."),
            ("", "2. Il presente contratto si applica esclusivamente a rapporti B2B (Business-to-Business). Il Cliente dichiara e garantisce di agire nell'ambito della propria attività professionale (Art. 71-bis disp. att. c.c.) e che non si applicano le disposizioni del Codice del Consumo (D.Lgs. 206/2005).")
        ]),
        ("3. NATURA DELLE FUNZIONALITÀ AI E CLAUSOLA DI MANLEVA CONTABILE E CIVILE", [
            ("3.1 Mero Strumento di Supporto Assistito (Human-in-the-loop)", [
                "1. Il Cliente riconosce ed accetta espressamente che le Funzionalità AI (OCR fatture, riconciliazione bancaria automatica, lettura moduli catastali, assistente chatbot) costituiscono uno strumento informatico di mero supporto assistito e pre-elaborazione dati.",
                "2. L'Intelligenza Artificiale non possiede personalità giuridica, non sostituisce il giudizio professionale dell'Amministratore e non assume alcuna responsabilità contabile, fiscale o civile.",
                "3. Il Cliente riconosce che i modelli di Intelligenza Artificiale, per loro natura tecnologica, possono generare imprecisioni, errate letture OCR, errate attribuzioni di cifre/date o cosiddette 'allucinazioni'."
            ]),
            ("3.2 Esclusiva Responsabilità dell'Amministratore (Art. 1130 e 1130-bis c.c.)", [
                "1. Ai sensi e per gli effetti degli Articoli 1130 e 1130-bis del Codice Civile italiano, l'Amministratore di Condominio rimane l'unico ed esclusivo responsabile legale e contabile per:\n   a) La tenuta e la veridicità del registro di anagrafe condominiale, del registro di contabilità e del registro dei verbali.\n   b) La correttezza aritmetica dei bilanci consuntivi, dei preventivi e dello stato patrimoniale.\n   c) L'esattezza delle ripartizioni delle spese per tabelle millesimali o criteri di legge/regolamento.\n   d) Il calcolo, la trattenuta ed il tempestivo versamento delle Ritenute d'Acconto (Modello F24 - 4%) all'Agenzia delle Entrate.\n   e) La validità delle lettere di sollecito e la gestione della morosità.",
                "2. Obbligo di Verifica Umana: Il Cliente ha l'obbligo inderogabile di verificare, controllare ed approvare manualmente ed aritmeticamente ogni singolo dato proposto o estratto dalle Funzionalità AI prima di salvarlo definitivamente sul database, inserirlo in un rendiconto o notificarlo ai condòmini."
            ]),
            ("3.3 Clausola di Manleva Totale da Impugnazioni e Sanzioni", [
                "1. Il Cliente manleva e tiene totalmente indenne M PROJECT SRL, i suoi legali rappresentanti, dipendenti e fornitori di infrastruttura da qualsiasi pretesa risarcitoria, azione legale, sanzione amministrativa/fiscale o danno derivante da:\n   a) Impugnazione di delibere assembleari ex Art. 1137 c.c. per presunti errori contabili, ripartizioni errate o vizi nei bilanci consuntivi e preventivi.\n   b) Sanzioni, cartelle esattoriali o ravvedimenti comminati dall'Agenzia delle Entrate per errato o tardivo versamento delle Ritenute d'Acconto (F24) o per errori nelle Certificazioni Uniche (CU) e 770.\n   c) Contestazioni da parte di condòmini o terzi fornitori relative a solleciti di pagamento errati, avvisi o comunicazioni inviate tramite la Piattaforma.\n   d) Perdita di dati derivante da errata cancellazione o modifica effettuata dall'Utente."
            ])
        ]),
        ("4. LIMITAZIONE DI RESPONSABILITÀ DEL FORNITORE", [
            ("", "1. Salvo i casi di dolo o colpa grave espressamente accertati con sentenza passata in giudicato, la responsabilità complessiva di M PROJECT SRL per qualsiasi malfunzionamento, interruzione del servizio, ritardo o errore tecnico della Piattaforma non potrà in alcun caso superare l'importo totale effettivamente pagato dal Cliente nei 12 mesi precedenti l'evento dannoso."),
            ("", "2. M PROJECT SRL non sarà in alcun caso responsabile per danni indiretti, lucri cessanti, perdita di opportunità commerciali, danno all'immagine dello studio professionale o sospensione dell'attività d'impresa.")
        ]),
        ("5. PROTEZIONE DEI DATI PERSONALI E NOMINA A RESPONSABILE DEL TRATTAMENTO (DPA ART. 28 GDPR)", [
            ("5.1 Ruoli di Privacy", [
                "1. Per quanto riguarda il trattamento dei dati personali dei condòmini, residenti e fornitori inseriti nella Piattaforma:\n   - Il Condominio / Cliente agisce in qualità di Titolare del Trattamento (Data Controller) ex Art. 4.7 GDPR.\n   - M PROJECT SRL agisce in qualità di Responsabile del Trattamento (Data Processor) ex Art. 28 GDPR."
            ]),
            ("5.2 Obblighi del Responsabile del Trattamento", [
                "1. M PROJECT SRL si impegna a trattare i dati personali al solo fine di erogare il Servizio e secondo le istruzioni documentate del Titolare.",
                "2. M PROJECT SRL adotta misure tecniche ed organizzative adeguate a garantire un livello di sicurezza commisurato al rischio (crittografia SSL/TLS, Row-Level Security nel database, backup regolari)."
            ]),
            ("5.3 Autorizzazione ai Sub-Responsabili (Sub-Processors)", [
                "1. Il Cliente autorizza espressamente M PROJECT SRL ad avvalersi dei seguenti sub-responsabili infrastrutturali per l'erogazione del servizio:\n   a) Supabase Inc. (Hosting Database PostgreSQL, Storage crittografato ed Edge Functions nell'UE/USA con clausole contrattuali tipo).\n   b) Google Cloud / Anthropic (Fornitori di modelli AI per OCR ed elaborazione testi con garanzia di non-utilizzo dei dati cliente per l'addestramento dei modelli pubblici).\n   c) Stripe Payments Europe Ltd. (Gestione dei pagamenti ricorsivi del SaaS).\n   d) Resend Inc. (Infrastruttura per l'invio telematico delle notifiche ed email di servizio)."
            ])
        ]),
        ("6. PROPRIETÀ DEI DATI E PORTABILITÀ", [
            ("", "1. Tutti i dati contabili, anagrafici, fatture e documenti inseriti dal Cliente rimangono di esclusiva proprietà del Condominio e del Cliente."),
            ("", "2. In qualsiasi momento durante la vigenza del contratto, ed entro 30 giorni dalla cessazione dell'abbonamento, il Cliente ha il diritto di esportare i propri dati in formato strutturato (Excel/CSV/PDF) tramite le funzionalità di export presenti sulla Piattaforma."),
            ("", "3. Decorso il termine di 30 giorni dalla cessazione dell'abbonamento o dalla scadenza della Prova Gratuita senza rinnovo, M PROJECT SRL si riserva la facoltà di cancellare o anonimizzare i dati dal database di produzione ai sensi delle politiche di ritenzione e minimizzazione del GDPR.")
        ]),
        ("7. CORRISPETTIVI, FATTURAZIONE E PROVA GRATUITA", [
            ("", "1. Prova Gratuita (Trial): Durante il periodo di Prova Gratuita (14 giorni), il Cliente può accedere alla Piattaforma a titolo gratuito. Alla scadenza del Trial, in assenza di selezione di un piano a pagamento, l'accesso verrà limitato in sola lettura o sospeso."),
            ("", "2. Abbonamenti a Pagamento: I corrispettivi per i Piani Base, Studio e Professional sono addebitati in via anticipata con cadenza mensile o annuale tramite carta di credito su circuito Stripe."),
            ("", "3. In caso di mancato o ritardato pagamento, M PROJECT SRL invierà un avviso e, decorsi 7 giorni, potrà sospendere l'accesso alla Piattaforma fino al saldo del dovuto.")
        ]),
        ("8. SOSPENSIONE E RISOLUZIONE DEL CONTRATTO", [
            ("", "1. M PROJECT SRL si riserva il diritto di sospendere o risolvere immediatamente il presente contratto in caso di:\n   a) Utilizzo della Piattaforma per attività illecite, invio di SPAM o email non autorizzate.\n   b) Tentativi di reverse-engineering, attacchi informatici o sovraccarico intenzionale delle API (es. abuso dei limiti di chiamata AI).")
        ]),
        ("9. LEGGE APPLICABILE E FORO COMPETENTE", [
            ("", "1. Il presente contratto è regolato esclusivamente dalla Legge Italiana."),
            ("", "2. Per qualsiasi controversia inerente l'interpretazione, esecuzione o risoluzione del presente contratto, le parti concordano che il Foro di [INSERIRE CITTA' FORO COMPETENTE, ES. MILANO / ROMA] sarà l'unico ed esclusivo Foro Competente.")
        ])
    ]

    for title, items in sections_content:
        doc.add_paragraph() # Spaziatore
        h = doc.add_heading(level=1)
        run_h = h.add_run(title)
        run_h.font.name = 'Arial'
        run_h.font.size = Pt(13)
        run_h.font.bold = True
        run_h.font.color.rgb = RGBColor(27, 54, 93)

        for prefix, body in items:
            if isinstance(body, list):
                sub_h = doc.add_heading(level=2)
                run_sub_h = sub_h.add_run(prefix)
                run_sub_h.font.name = 'Arial'
                run_sub_h.font.size = Pt(11)
                run_sub_h.font.bold = True
                run_sub_h.font.color.rgb = RGBColor(59, 130, 246)

                for line in body:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.2)
                    run_line = p.add_run(line)
                    run_line.font.name = 'Arial'
                    run_line.font.size = Pt(10)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.1)
                if prefix:
                    run_p = p.add_run(prefix)
                    run_p.font.name = 'Arial'
                    run_p.font.size = Pt(10)
                    run_p.font.bold = True
                run_b = p.add_run(body)
                run_b.font.name = 'Arial'
                run_b.font.size = Pt(10)

    output_path = "/Users/gabrielemaesani/Documents/CondoAI2/TERMINI_E_CONDIZIONI_M_PROJECT_SRL.docx"
    doc.save(output_path)
    print(f"File Word generato con successo: {output_path}")

if __name__ == '__main__':
    create_tos_docx()
